import io
from django.http import HttpResponse
from django.conf import settings


class ExcelExporter:
    """Export data to Excel format."""
    
    @staticmethod
    def export_to_excel(data, filename, sheet_name='Sheet1'):
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = sheet_name
        
        if not data:
            return wb
        
        # Header styling
        header_font = Font(bold=True, color='FFFFFF')
        header_fill = PatternFill(start_color='1A237E', end_color='1A237E', fill_type='solid')
        header_alignment = Alignment(horizontal='center', vertical='center')
        thin_border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )
        
        # Write headers
        headers = list(data[0].keys()) if data else []
        for col_num, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col_num, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_alignment
            cell.border = thin_border
        
        # Write data
        for row_num, row_data in enumerate(data, 2):
            for col_num, header in enumerate(headers, 1):
                cell = ws.cell(row=row_num, column=col_num, value=row_data.get(header, ''))
                cell.border = thin_border
                cell.alignment = Alignment(horizontal='center', vertical='center')
        
        # Auto-adjust column widths
        for col in ws.columns:
            max_length = 0
            col_letter = col[0].column_letter
            for cell in col:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            ws.column_dimensions[col_letter].width = min(max_length + 2, 50)
        
        return wb


class PDFExporter:
    """Export data to PDF format."""
    
    @staticmethod
    def export_to_pdf(data, filename, title='Report'):
        from reportlab.lib.pagesizes import A4, letter
        from reportlab.lib import colors
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
        from reportlab.lib.units import inch
        
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4)
        elements = []
        
        # Title
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            spaceAfter=30,
            textColor=colors.HexColor('#1A237E')
        )
        elements.append(Paragraph(title, title_style))
        elements.append(Spacer(1, 12))
        
        if data:
            headers = list(data[0].keys())
            table_data = [headers]
            for row in data:
                table_data.append([str(row.get(h, '')) for h in headers])
            
            table = Table(table_data)
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1A237E')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F5F5F5')]),
            ]))
            elements.append(table)
        
        doc.build(elements)
        buffer.seek(0)
        return buffer


class WordExporter:
    """Export data to Word format."""
    
    @staticmethod
    def export_to_word(data, filename, title='Report'):
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        
        doc = Document()
        
        # Title
        title_para = doc.add_heading(title, 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title_para.runs:
            run.font.color.rgb = RGBColor(26, 35, 126)  # #1A237E
        
        doc.add_paragraph()
        
        if data:
            headers = list(data[0].keys())
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Header row
            hdr_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
                for paragraph in hdr_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
            
            # Data rows
            for row_data in data:
                row_cells = table.add_row().cells
                for i, header in enumerate(headers):
                    row_cells[i].text = str(row_data.get(header, ''))
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer


class CSVExporter:
    """Export data to CSV format."""
    
    @staticmethod
    def export_to_csv(data, filename):
        import csv
        
        output = io.StringIO()
        if data:
            headers = list(data[0].keys())
            writer = csv.DictWriter(output, fieldnames=headers)
            writer.writeheader()
            writer.writerows(data)
        
        response = HttpResponse(output.getvalue(), content_type='text/csv')
        response['Content-Disposition'] = f'attachment; filename="{filename}.csv"'
        return response


class JPEGExporter:
    """Export data to JPEG format."""
    
    @staticmethod
    def export_to_jpeg(data, filename, title='Report'):
        from PIL import Image, ImageDraw, ImageFont
        
        # Create image
        width = 1200
        row_height = 30
        header_height = 50
        margin = 50
        
        if data:
            num_rows = len(data) + 1  # +1 for header
        else:
            num_rows = 1
        
        height = margin * 2 + header_height + (num_rows * row_height)
        
        img = Image.new('RGB', (width, height), color='white')
        draw = ImageDraw.Draw(img)
        
        # Title
        try:
            title_font = ImageFont.truetype("arial.ttf", 24)
            header_font = ImageFont.truetype("arial.ttf", 14)
            data_font = ImageFont.truetype("arial.ttf", 12)
        except:
            title_font = ImageFont.load_default()
            header_font = ImageFont.load_default()
            data_font = ImageFont.load_default()
        
        draw.text((margin, margin), title, fill='#1A237E', font=title_font)
        
        if data:
            headers = list(data[0].keys())
            col_width = (width - margin * 2) // len(headers)
            
            # Draw headers
            y = margin + 50
            draw.rectangle([margin, y, width - margin, y + header_height], fill='#1A237E')
            for i, header in enumerate(headers):
                x = margin + (i * col_width)
                draw.text((x + 10, y + 15), header, fill='white', font=header_font)
            
            # Draw data rows
            y += header_height
            for row_num, row_data in enumerate(data):
                bg_color = '#F5F5F5' if row_num % 2 == 0 else 'white'
                draw.rectangle([margin, y, width - margin, y + row_height], fill=bg_color, outline='#CCCCCC')
                for i, header in enumerate(headers):
                    x = margin + (i * col_width)
                    draw.text((x + 10, y + 8), str(row_data.get(header, '')), fill='black', font=data_font)
                y += row_height
        
        buffer = io.BytesIO()
        img.save(buffer, format='JPEG', quality=95)
        buffer.seek(0)
        return buffer
