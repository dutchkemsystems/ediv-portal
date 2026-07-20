import io
from django.core.files.uploadedfile import InMemoryUploadedFile


class ExcelImporter:
    """Import data from Excel format."""
    
    @staticmethod
    def import_from_excel(file):
        import openpyxl
        
        wb = openpyxl.load_workbook(file)
        ws = wb.active
        
        data = []
        headers = []
        
        for row_num, row in enumerate(ws.iter_rows(values_only=True)):
            if row_num == 0:
                headers = [str(h).strip() if h else f'Column_{i}' for i, h in enumerate(row)]
            else:
                row_dict = {}
                for i, value in enumerate(row):
                    if i < len(headers):
                        row_dict[headers[i]] = value
                data.append(row_dict)
        
        return {
            'headers': headers,
            'data': data,
            'total_rows': len(data)
        }


class CSVImporter:
    """Import data from CSV format."""
    
    @staticmethod
    def import_from_csv(file):
        import csv
        
        # Decode file content
        if isinstance(file, InMemoryUploadedFile):
            content = file.read().decode('utf-8')
        else:
            content = file.read().decode('utf-8')
        
        reader = csv.DictReader(io.StringIO(content))
        
        data = []
        headers = reader.fieldnames or []
        
        for row in reader:
            data.append(dict(row))
        
        return {
            'headers': headers,
            'data': data,
            'total_rows': len(data)
        }


class PDFImporter:
    """Import/extract data from PDF format."""
    
    @staticmethod
    def extract_text_from_pdf(file):
        from PyPDF2 import PdfReader
        
        reader = PdfReader(file)
        text_content = []
        
        for page in reader.pages:
            text_content.append(page.extract_text())
        
        return {
            'pages': len(reader.pages),
            'content': '\n\n'.join(text_content)
        }


class WordImporter:
    """Import/extract data from Word format."""
    
    @staticmethod
    def extract_from_word(file):
        from docx import Document
        
        doc = Document(file)
        
        content = []
        tables = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                content.append(para.text)
        
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            tables.append(table_data)
        
        return {
            'paragraphs': content,
            'tables': tables,
            'total_paragraphs': len(content),
            'total_tables': len(tables)
        }


class DataValidator:
    """Validate imported data."""
    
    @staticmethod
    def validate_student_data(data):
        errors = []
        valid_data = []
        
        required_fields = ['first_name', 'last_name', 'email', 'admission_number']
        
        for row_num, row in enumerate(data, 1):
            row_errors = []
            
            for field in required_fields:
                if not row.get(field):
                    row_errors.append(f'{field} is required')
            
            if row.get('email') and not '@' in str(row['email']):
                row_errors.append('Invalid email format')
            
            if row_errors:
                errors.append({'row': row_num, 'errors': row_errors})
            else:
                valid_data.append(row)
        
        return {
            'valid': len(errors) == 0,
            'total_rows': len(data),
            'valid_rows': len(valid_data),
            'error_rows': len(errors),
            'errors': errors,
            'data': valid_data
        }
    
    @staticmethod
    def validate_staff_data(data):
        errors = []
        valid_data = []
        
        required_fields = ['first_name', 'last_name', 'email', 'staff_id', 'employee_number']
        
        for row_num, row in enumerate(data, 1):
            row_errors = []
            
            for field in required_fields:
                if not row.get(field):
                    row_errors.append(f'{field} is required')
            
            if row_errors:
                errors.append({'row': row_num, 'errors': row_errors})
            else:
                valid_data.append(row)
        
        return {
            'valid': len(errors) == 0,
            'total_rows': len(data),
            'valid_rows': len(valid_data),
            'error_rows': len(errors),
            'errors': errors,
            'data': valid_data
        }
