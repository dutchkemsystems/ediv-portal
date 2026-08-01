"""Import/export service for file data operations."""
import io
import csv
import datetime
import traceback
from django.core.files.base import ContentFile
from django.utils import timezone
from apps.files.models import File, FileAttachment, FileMovement


class ImportExportService:
    """Service for importing and exporting files in multiple formats."""

    SUPPORTED_IMPORT_FORMATS = ['doc', 'docx', 'xls', 'xlsx', 'pdf', 'jpeg', 'png', 'csv', 'txt']
    SUPPORTED_EXPORT_FORMATS = ['xlsx', 'csv', 'pdf', 'docx']

    @staticmethod
    def import_file(*, uploaded_file, file_format, created_by,
                    department=None, default_classification='INTERNAL',
                    default_priority='NORMAL') -> dict:
        """
        Import a document and create a File record.

        Returns: {'file': File, 'attachments': [FileAttachment], 'errors': []}
        """
        errors = []
        attachments = []
        file_obj = None

        if file_format not in ImportExportService.SUPPORTED_IMPORT_FORMATS:
            return {'file': None, 'attachments': [], 'errors': [f'Unsupported format: {file_format}']}

        try:
            # Read file content for description extraction
            description = ''
            metadata = {}

            if file_format == 'txt':
                uploaded_file.seek(0)
                content = uploaded_file.read()
                try:
                    description = content.decode('utf-8')
                except UnicodeDecodeError:
                    description = content.decode('latin-1')

            elif file_format == 'csv':
                uploaded_file.seek(0)
                raw = uploaded_file.read()
                try:
                    text = raw.decode('utf-8')
                except UnicodeDecodeError:
                    text = raw.decode('latin-1')
                reader = csv.DictReader(io.StringIO(text))
                rows = list(reader)
                if rows:
                    first_row = rows[0]
                    # Use 'title' column if present
                    if 'title' in first_row:
                        metadata['title'] = first_row['title']
                    if 'description' in first_row:
                        description = first_row['description']
                    if 'file_type' in first_row:
                        metadata['file_type'] = first_row['file_type']
                    if 'priority' in first_row:
                        metadata['priority'] = first_row['priority']
                # Re-encode for attachment
                uploaded_file = ContentFile(raw, name=getattr(uploaded_file, 'name', 'import.csv'))

            elif file_format in ('xlsx', 'xls'):
                uploaded_file.seek(0)
                metadata = ImportExportService._import_spreadsheet(uploaded_file, file_format)
                if metadata.get('data'):
                    d = metadata['data']
                    if 'title' in d:
                        metadata['title'] = d['title']
                    if 'description' in d:
                        description = d['description']

            elif file_format in ('docx', 'pdf'):
                uploaded_file.seek(0)
                try:
                    doc_data = ImportExportService._import_document(uploaded_file, file_format)
                    description = doc_data.get('text', '')
                except ImportError as e:
                    errors.append(f'Missing library for {file_format}: {str(e)}')

            elif file_format in ('jpeg', 'png'):
                uploaded_file.seek(0)
                try:
                    img_data = ImportExportService._import_image(uploaded_file, file_format)
                    description = f"Image: {img_data.get('width', 0)}x{img_data.get('height', 0)}"
                except ImportError as e:
                    errors.append(f'Missing library for image processing: {str(e)}')

            # Determine title
            file_name = getattr(uploaded_file, 'name', f'import.{file_format}')
            title = metadata.get('title', file_name.rsplit('.', 1)[0] if '.' in file_name else file_name)

            # Determine file_type
            file_type = metadata.get('file_type', 'OTHER')
            valid_types = [c[0] for c in File._meta.get_field('file_type').choices]
            if file_type not in valid_types:
                file_type = 'OTHER'

            # Determine priority
            priority = metadata.get('priority', default_priority)
            valid_priorities = ['LOW', 'NORMAL', 'HIGH', 'URGENT']
            if priority not in valid_priorities:
                priority = default_priority

            # Create File record
            file_obj = File.objects.create(
                file_number=ImportExportService._generate_file_number(),
                title=title,
                file_type=file_type,
                file_category='ADMIN',
                description=description,
                created_by=created_by,
                current_holder=created_by,
                department=department,
                classification=default_classification,
                priority=priority,
                status='DRAFT',
                tags=[],
            )

            # Create attachment
            uploaded_file.seek(0)
            attachment = FileAttachment.objects.create(
                file=file_obj,
                document=uploaded_file,
                original_filename=file_name,
                file_size=uploaded_file.size if hasattr(uploaded_file, 'size') else 0,
                mime_type='',
                file_format=file_format,
                uploaded_by=created_by,
            )
            attachments.append(attachment)

            # Record CREATED movement
            FileMovement.objects.create(
                file=file_obj,
                from_holder=created_by,
                action=FileMovement.Action.CREATED,
                remarks=f'File imported from {file_format.upper()} document',
            )

        except Exception as e:
            errors.append(f'Import error: {str(e)}')

        return {'file': file_obj, 'attachments': attachments, 'errors': errors}

    @staticmethod
    def _generate_file_number():
        """Generate a unique file number."""
        year = datetime.date.today().year
        seq = File.objects.filter(file_number__startswith=f'EDIV-{year}-IMP').count() + 1
        return f'EDIV-{year}-IMP-{seq:04d}'

    @staticmethod
    def _import_spreadsheet(uploaded_file, file_format) -> dict:
        """Import Excel/CSV spreadsheet. Returns extracted metadata."""
        if file_format == 'xlsx':
            import openpyxl
            wb = openpyxl.load_workbook(uploaded_file, read_only=True, data_only=True)
            ws = wb.active
            headers = []
            data = {}
            for row_idx, row in enumerate(ws.iter_rows(values_only=True), 1):
                if row_idx == 1:
                    headers = [str(cell) for cell in row if cell is not None]
                elif row_idx == 2:
                    for col_idx, header in enumerate(headers):
                        if col_idx < len(row) and row[col_idx] is not None:
                            data[header.lower().strip()] = str(row[col_idx])
                    break
            wb.close()
            return {'headers': headers, 'data': data, 'row_count': max(ws.max_row - 1, 0)}

        elif file_format == 'csv':
            uploaded_file.seek(0)
            raw = uploaded_file.read()
            text = raw.decode('utf-8')
            reader = csv.DictReader(io.StringIO(text))
            rows = list(reader)
            return {
                'headers': list(rows[0].keys()) if rows else [],
                'data': rows[0] if rows else {},
                'row_count': len(rows)
            }
        return {}

    @staticmethod
    def _import_document(uploaded_file, file_format) -> dict:
        """Import Word/PDF document. Returns extracted text."""
        if file_format == 'docx':
            from docx import Document
            doc = Document(uploaded_file)
            text = '\n'.join([para.text for para in doc.paragraphs if para.text])
            return {'text': text, 'paragraph_count': len(doc.paragraphs)}

        elif file_format == 'pdf':
            import PyPDF2
            reader = PyPDF2.PdfReader(uploaded_file)
            text = ''
            for page in reader.pages:
                text += page.extract_text() or ''
            return {'text': text, 'page_count': len(reader.pages)}

        return {'text': '', 'page_count': 0}

    @staticmethod
    def _import_image(uploaded_file, file_format) -> dict:
        """Import image file. Returns metadata."""
        from PIL import Image
        img = Image.open(uploaded_file)
        return {
            'width': img.width,
            'height': img.height,
            'format': img.format,
            'mode': img.mode,
        }

    @staticmethod
    def export_files(*, file_ids, export_format, exported_by) -> ContentFile:
        """
        Export files to specified format.

        Returns ContentFile with exported content.
        Returns None for unsupported format.
        """
        if export_format not in ImportExportService.SUPPORTED_EXPORT_FORMATS:
            return None

        # Fetch files (skip non-existent IDs)
        files = File.objects.filter(id__in=file_ids).select_related(
            'created_by', 'department'
        )

        if export_format == 'xlsx':
            return ImportExportService._export_to_xlsx(files)
        elif export_format == 'csv':
            return ImportExportService._export_to_csv(files)
        elif export_format == 'pdf':
            return ImportExportService._export_to_pdf(files)
        elif export_format == 'docx':
            return ImportExportService._export_to_docx(files)

        return None

    @staticmethod
    def _export_to_xlsx(files) -> ContentFile:
        """Export files to Excel."""
        import openpyxl
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Files Export"

        headers = ['File Number', 'Title', 'Type', 'Category', 'Status',
                   'Classification', 'Priority', 'Created By', 'Department',
                   'Created At', 'Due Date', 'Tags']
        ws.append(headers)

        for f in files:
            ws.append([
                f.file_number,
                f.title,
                f.get_file_type_display(),
                f.get_file_category_display(),
                f.get_status_display(),
                f.get_classification_display(),
                f.get_priority_display(),
                f.created_by.get_full_name() or f.created_by.username,
                f.department.name if f.department else '',
                f.created_at.strftime('%Y-%m-%d %H:%M') if f.created_at else '',
                f.due_date.strftime('%Y-%m-%d') if f.due_date else '',
                ', '.join(f.tags) if f.tags else '',
            ])

        buffer = io.BytesIO()
        wb.save(buffer)
        buffer.seek(0)
        return ContentFile(buffer.getvalue(), name=f'files_export_{timezone.now().strftime("%Y%m%d_%H%M%S")}.xlsx')

    @staticmethod
    def _export_to_csv(files) -> ContentFile:
        """Export files to CSV."""
        output = io.StringIO()
        writer = csv.writer(output)

        headers = ['File Number', 'Title', 'Type', 'Category', 'Status',
                   'Classification', 'Priority', 'Created By', 'Department',
                   'Created At', 'Due Date']
        writer.writerow(headers)

        for f in files:
            writer.writerow([
                f.file_number,
                f.title,
                f.file_type,
                f.file_category,
                f.status,
                f.classification,
                f.priority,
                f.created_by.get_full_name() or f.created_by.username,
                f.department.name if f.department else '',
                f.created_at.strftime('%Y-%m-%d %H:%M') if f.created_at else '',
                f.due_date.strftime('%Y-%m-%d') if f.due_date else '',
            ])

        content = output.getvalue().encode('utf-8')
        return ContentFile(content, name=f'files_export_{timezone.now().strftime("%Y%m%d_%H%M%S")}.csv')

    @staticmethod
    def _export_to_pdf(files) -> ContentFile:
        """Export files to PDF using reportlab. Returns None if reportlab is not installed."""
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.lib import colors
            from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet
        except ImportError:
            return None

        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4)
        elements = []

        styles = getSampleStyleSheet()
        elements.append(Paragraph("Files Export", styles['Title']))
        elements.append(Spacer(1, 20))

        data = [['File Number', 'Title', 'Type', 'Status', 'Priority', 'Created By']]
        for f in files:
            data.append([
                f.file_number,
                f.title[:40] + '...' if len(f.title) > 40 else f.title,
                f.file_type,
                f.status,
                f.priority,
                f.created_by.get_full_name() or f.created_by.username,
            ])

        table = Table(data, colWidths=[80, 150, 70, 70, 60, 100])
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTSIZE', (0, 0), (-1, 0), 10),
            ('FONTSIZE', (0, 1), (-1, -1), 8),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.Color(0.95, 0.95, 0.95)]),
        ]))
        elements.append(table)

        doc.build(elements)
        buffer.seek(0)
        return ContentFile(buffer.getvalue(), name=f'files_export_{timezone.now().strftime("%Y%m%d_%H%M%S")}.pdf')

    @staticmethod
    def _export_to_docx(files) -> ContentFile:
        """Export files to Word document."""
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.table import WD_TABLE_ALIGNMENT

        doc = Document()
        doc.add_heading('Files Export', 0)

        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        headers = ['File Number', 'Title', 'Type', 'Status', 'Priority', 'Created By']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        for f in files:
            row = table.add_row()
            row.cells[0].text = f.file_number
            row.cells[1].text = f.title[:50]
            row.cells[2].text = f.file_type
            row.cells[3].text = f.status
            row.cells[4].text = f.priority
            row.cells[5].text = f.created_by.get_full_name() or f.created_by.username

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return ContentFile(buffer.getvalue(), name=f'files_export_{timezone.now().strftime("%Y%m%d_%H%M%S")}.docx')

    @staticmethod
    def bulk_import(*, uploaded_file, file_format, created_by,
                    department=None, skip_errors=True) -> dict:
        """
        Bulk import from CSV/XLSX.
        Returns: {'imported': int, 'skipped': int, 'errors': list}
        """
        imported = 0
        skipped = 0
        errors = []

        if file_format not in ('csv', 'xlsx'):
            return {'imported': 0, 'skipped': 0, 'errors': [f'Bulk import only supports csv/xlsx, got: {file_format}']}

        try:
            if file_format == 'csv':
                uploaded_file.seek(0)
                raw = uploaded_file.read()
                text = raw.decode('utf-8')
                reader = csv.DictReader(io.StringIO(text))
                rows = list(reader)
            elif file_format == 'xlsx':
                import openpyxl
                uploaded_file.seek(0)
                wb = openpyxl.load_workbook(uploaded_file, read_only=True, data_only=True)
                ws = wb.active
                headers = None
                rows = []
                for row_idx, row in enumerate(ws.iter_rows(values_only=True), 1):
                    if row_idx == 1:
                        headers = [str(cell) if cell else f'col_{i}' for i, cell in enumerate(row)]
                    else:
                        row_dict = {}
                        for col_idx, cell in enumerate(row):
                            if col_idx < len(headers):
                                row_dict[headers[col_idx]] = str(cell) if cell is not None else ''
                        rows.append(row_dict)
                wb.close()
            else:
                rows = []

            for idx, row in enumerate(rows):
                try:
                    title = row.get('title', '') or row.get('Title', '') or f'Bulk Import {idx + 1}'
                    file_type = row.get('file_type', '') or row.get('Type', '') or 'OTHER'

                    # Validate file_type
                    valid_types = [c[0] for c in File._meta.get_field('file_type').choices]
                    if file_type not in valid_types:
                        file_type = 'OTHER'

                    priority = row.get('priority', '') or row.get('Priority', '') or 'NORMAL'
                    valid_priorities = ['LOW', 'NORMAL', 'HIGH', 'URGENT']
                    if priority not in valid_priorities:
                        priority = 'NORMAL'

                    description = row.get('description', '') or row.get('Description', '')

                    file_obj = File.objects.create(
                        file_number=ImportExportService._generate_file_number(),
                        title=title,
                        file_type=file_type,
                        file_category='ADMIN',
                        description=description,
                        created_by=created_by,
                        current_holder=created_by,
                        department=department,
                        classification='INTERNAL',
                        priority=priority,
                        status='DRAFT',
                        tags=[],
                    )

                    FileMovement.objects.create(
                        file=file_obj,
                        from_holder=created_by,
                        action=FileMovement.Action.CREATED,
                        remarks='Bulk imported',
                    )

                    imported += 1

                except Exception as e:
                    if skip_errors:
                        skipped += 1
                        errors.append(f'Row {idx + 1}: {str(e)}')
                    else:
                        errors.append(f'Row {idx + 1}: {str(e)}')
                        break

        except Exception as e:
            errors.append(f'Bulk import error: {str(e)}')

        return {'imported': imported, 'skipped': skipped, 'errors': errors}
