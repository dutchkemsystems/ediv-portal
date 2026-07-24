import csv
import io
import json
from datetime import datetime

from django.http import HttpResponse
from django.db import models
from rest_framework import viewsets, permissions, status
from rest_framework.decorators import action
from rest_framework.response import Response

from config.security import AuditLogger
from .models import ImportJob, ImportError as ImportErrorModel
from .serializers import ImportJobSerializer, ImportJobListSerializer


EXPORTABLE_MODELS = {
    'students': {
        'model_path': 'apps.students.models',
        'model_name': 'Student',
        'fields': ['id', 'first_name', 'last_name', 'email', 'student_id', 'date_of_birth', 'gender'],
    },
    'staff': {
        'model_path': 'apps.staff.models',
        'model_name': 'Staff',
        'fields': ['id', 'first_name', 'last_name', 'email', 'employee_id', 'department', 'position'],
    },
    'schools': {
        'model_path': 'apps.schools.models',
        'model_name': 'School',
        'fields': ['id', 'name', 'code', 'address', 'phone', 'email', 'principal_name'],
    },
}


def _import_rows(rows, target_model, job):
    imported = 0
    errors = []
    model_config = EXPORTABLE_MODELS.get(target_model)
    if not model_config:
        return 0, [{'row': 0, 'field': '', 'error': f'Unknown model: {target_model}', 'raw': ''}]

    try:
        parts = model_config['model_path'].split('.')
        module = __import__(parts[0], fromlist=[parts[1]])
        for part in parts[1:]:
            module = getattr(module, part)
        model_cls = getattr(module, model_config['model_name'])
    except (ImportError, AttributeError) as e:
        return 0, [{'row': 0, 'field': '', 'error': f'Model load failed: {e}', 'raw': ''}]

    valid_fields = set(f.name for f in model_cls._meta.get_fields() if hasattr(f, 'column'))

    for i, row in enumerate(rows, start=1):
        try:
            filtered = {k: v for k, v in row.items() if k in valid_fields and v}
            if not filtered:
                errors.append({'row': i, 'field': '', 'error': 'No valid fields', 'raw': json.dumps(row)})
                continue
            model_cls.objects.create(**filtered)
            imported += 1
        except Exception as e:
            errors.append({'row': i, 'field': '', 'error': str(e), 'raw': json.dumps(row)})

    return imported, errors


class ImportJobViewSet(viewsets.ModelViewSet):
    permission_classes = [permissions.IsAuthenticated]

    def get_serializer_class(self):
        if self.action == 'list':
            return ImportJobListSerializer
        return ImportJobSerializer

    def get_queryset(self):
        user = self.request.user
        if user.role in ('SYSADMIN', 'TG', 'PS'):
            return ImportJob.objects.all()
        return ImportJob.objects.filter(created_by=user)

    def perform_create(self, serializer):
        serializer.save(created_by=self.request.user)

    @action(detail=False, methods=['post'], url_path='import')
    def import_data(self, request):
        uploaded_file = request.FILES.get('file')
        target_model = request.data.get('model', '')

        if not uploaded_file:
            return Response({'error': 'No file provided.'}, status=status.HTTP_400_BAD_REQUEST)
        if target_model not in EXPORTABLE_MODELS:
            return Response({'error': f'Invalid model. Choose from: {list(EXPORTABLE_MODELS.keys())}'},
                            status=status.HTTP_400_BAD_REQUEST)

        file_name = uploaded_file.name
        ext = file_name.rsplit('.', 1)[-1].lower() if '.' in file_name else ''

        type_map = {'csv': 'CSV', 'xlsx': 'EXCEL', 'xls': 'EXCEL', 'pdf': 'PDF',
                     'docx': 'WORD', 'json': 'JSON'}
        file_type = type_map.get(ext, 'CSV')

        job = ImportJob.objects.create(
            file_name=file_name,
            file_type=file_type,
            target_model=target_model,
            status='PROCESSING',
            created_by=request.user,
        )

        rows = []
        try:
            if ext == 'csv':
                content = uploaded_file.read().decode('utf-8-sig')
                reader = csv.DictReader(io.StringIO(content))
                rows = list(reader)
            elif ext in ('xlsx', 'xls'):
                import openpyxl
                wb = openpyxl.load_workbook(uploaded_file, read_only=True)
                ws = wb.active
                headers = [str(cell.value).strip() if cell.value else '' for cell in next(ws.iter_rows(max_row=1))]
                for row in ws.iter_rows(min_row=2, values_only=True):
                    rows.append({headers[i]: str(v) if v is not None else '' for i, v in enumerate(row) if i < len(headers)})
                wb.close()
            elif ext == 'json':
                content = uploaded_file.read().decode('utf-8')
                data = json.loads(content)
                if isinstance(data, list):
                    rows = data
                elif isinstance(data, dict) and 'data' in data:
                    rows = data['data']
                else:
                    rows = [data]
            elif ext == 'pdf':
                import pdfplumber
                with pdfplumber.open(uploaded_file) as pdf:
                    for page in pdf.pages:
                        table = page.extract_table()
                        if table and len(table) > 1:
                            headers = [str(h).strip() if h else '' for h in table[0]]
                            for row in table[1:]:
                                rows.append({headers[i]: str(v) if v else '' for i, v in enumerate(row) if i < len(headers)})
            elif ext == 'docx':
                import docx
                doc = docx.Document(uploaded_file)
                for table in doc.tables:
                    if len(table.rows) > 1:
                        headers = [cell.text.strip() for cell in table.rows[0].cells]
                        for row in table.rows[1:]:
                            rows.append({headers[i]: cell.text.strip() for i, cell in enumerate(row.cells) if i < len(headers)})
            else:
                job.status = 'FAILED'
                job.error_log = [{'row': 0, 'field': '', 'error': f'Unsupported file type: {ext}', 'raw': ''}]
                job.completed_at = datetime.now()
                job.save()
                return Response({'error': f'Unsupported file type: {ext}'}, status=status.HTTP_400_BAD_REQUEST)
        except Exception as e:
            job.status = 'FAILED'
            job.error_log = [{'row': 0, 'field': '', 'error': f'Parse error: {e}', 'raw': ''}]
            job.completed_at = datetime.now()
            job.save()
            return Response({'error': f'File parse error: {e}'}, status=status.HTTP_400_BAD_REQUEST)

        job.total_rows = len(rows)
        imported, import_errors = _import_rows(rows, target_model, job)

        for err in import_errors:
            ImportErrorModel.objects.create(
                job=job,
                row_number=err['row'],
                field_name=err.get('field', ''),
                error_message=err['error'],
                raw_value=err.get('raw', ''),
            )

        job.success_rows = imported
        job.error_rows = len(import_errors)
        job.status = 'COMPLETED' if not import_errors else ('FAILED' if imported == 0 else 'COMPLETED')
        job.error_log = import_errors
        job.completed_at = datetime.now()
        job.save()

        AuditLogger.log_action(
            user=request.user,
            action='IMPORT',
            resource_type='DataImport',
            resource_id=job.id,
            description=f"Imported {imported}/{job.total_rows} rows from {file_name} to {target_model}",
        )

        return Response(ImportJobSerializer(job).data, status=status.HTTP_201_CREATED)

    @action(detail=False, methods=['get'], url_path='export')
    def export_data(self, request):
        target_model = request.query_params.get('model', '')
        fmt = request.query_params.get('format', 'csv').lower()

        if target_model not in EXPORTABLE_MODELS:
            return Response({'error': f'Invalid model. Choose from: {list(EXPORTABLE_MODELS.keys())}'},
                            status=status.HTTP_400_BAD_REQUEST)

        config = EXPORTABLE_MODELS[target_model]
        try:
            parts = config['model_path'].split('.')
            module = __import__(parts[0], fromlist=[parts[1]])
            for part in parts[1:]:
                module = getattr(module, part)
            model_cls = getattr(module, config['model_name'])
        except (ImportError, AttributeError) as e:
            return Response({'error': f'Model load failed: {e}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

        qs = model_cls.objects.all()
        fields = [f for f in config['fields'] if f != 'id' or True]
        data_fields = [f for f in config['fields'] if f != 'id']

        if hasattr(model_cls, 'first_name') and hasattr(model_cls, 'last_name'):
            rows = list(qs.values(*data_fields))
        else:
            rows = list(qs.values(*data_fields))

        AuditLogger.log_action(
            user=request.user,
            action='EXPORT',
            resource_type='DataExport',
            description=f"Exported {len(rows)} {target_model} as {fmt.upper()}",
        )

        if fmt == 'csv':
            response = HttpResponse(content_type='text/csv')
            response['Content-Disposition'] = f'attachment; filename="{target_model}_export.csv"'
            writer = csv.DictWriter(response, fieldnames=data_fields)
            writer.writeheader()
            writer.writerows(rows)
            return response

        elif fmt == 'excel' or fmt == 'xlsx':
            import openpyxl
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = target_model.title()
            ws.append(data_fields)
            for row in rows:
                ws.append([str(row.get(f, '')) for f in data_fields])
            response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
            response['Content-Disposition'] = f'attachment; filename="{target_model}_export.xlsx"'
            wb.save(response)
            return response

        elif fmt == 'pdf':
            from reportlab.lib import colors
            from reportlab.lib.pagesizes import landscape, A4
            from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph
            from reportlab.lib.styles import getSampleStyleSheet

            response = HttpResponse(content_type='application/pdf')
            response['Content-Disposition'] = f'attachment; filename="{target_model}_export.pdf"'

            doc = SimpleDocTemplate(response, pagesize=landscape(A4))
            elements = []
            styles = getSampleStyleSheet()
            elements.append(Paragraph(f"{target_model.title()} Export", styles['Title']))

            table_data = [data_fields]
            for row in rows:
                table_data.append([str(row.get(f, '')) for f in data_fields])

            table = Table(table_data)
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1976d2')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('FONTSIZE', (0, 1), (-1, -1), 8),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f5f5f5')]),
            ]))
            elements.append(table)
            doc.build(elements)
            return response

        elif fmt == 'word' or fmt == 'docx':
            from docx import Document
            from docx.shared import Inches

            doc = Document()
            doc.add_heading(f'{target_model.title()} Export', 0)

            table = doc.add_table(rows=1, cols=len(data_fields))
            table.style = 'Light Grid Accent 1'
            for i, field in enumerate(data_fields):
                table.rows[0].cells[i].text = field

            for row in rows:
                cells = table.add_row().cells
                for i, field in enumerate(data_fields):
                    cells[i].text = str(row.get(field, ''))

            response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = f'attachment; filename="{target_model}_export.docx"'
            doc.save(response)
            return response

        elif fmt == 'json':
            data = [{k: str(v) if v is not None else '' for k, v in row.items()} for row in rows]
            response = HttpResponse(json.dumps(data, indent=2), content_type='application/json')
            response['Content-Disposition'] = f'attachment; filename="{target_model}_export.json"'
            return response

        else:
            return Response({'error': f'Unsupported format: {fmt}. Use csv, excel, pdf, word, or json.'},
                            status=status.HTTP_400_BAD_REQUEST)
